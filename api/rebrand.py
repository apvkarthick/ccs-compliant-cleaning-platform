"""SDS document rebranding: replace supplier identity with CCS identity."""
from __future__ import annotations

import io
import re
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# Per-brand header logo: smart_clean uses SmartClean logo; others keep original logo
_BRAND_LOGOS: dict[str, Path] = {
    "smart_clean": Path(__file__).parent / "assets" / "smartclean_logo.jpg",
}


CCS = {
    "supplier_name": "Compliant Cleaning Supplies & Systems PTY LTD",
    "address": "86 Crockford Street, NORTHGATE QLD 4013",
    "po_box": "PO Box 258, Hamilton Central QLD 4007",
    "telephone": "1300 314 491",
    "emergency": "131126",
    "email": "sales@compliantcs.com.au",
    "website": "https://www.compliantcs.com.au/",
}
_CCS_EMAIL_URL = "mailto:sales@compliantcs.com.au"
_CCS_WEB_URL = "https://www.compliantcs.com.au/"
_CCS_ABN = "27 144 521 200"

# Maps paragraph label prefixes to CCS field keys (mode: "full" or "tab_value")
_SUPPLIER_FIELDS: list[tuple[str, str, str]] = [
    ("Supplier Name", "supplier_name", "full"),
    ("Address", "address", "tab_value"),
    ("Telephone", "telephone", "tab_value"),
    ("Emergency", "emergency", "tab_value"),
    ("Email", "email", "tab_value"),
    ("Web Site", "website", "tab_value"),
    ("Website", "website", "tab_value"),
]

_CCS_DOMAIN = "compliantcs.com.au"


def rebrand_sds(docx_bytes: bytes, sds_date: str | None = None, brand: str = "") -> tuple[bytes, dict]:
    """
    Rebrand a supplier SDS DOCX with CCS identity.

    Returns (rebranded_bytes, summary_dict).
    sds_date: DD/MM/YYYY string; defaults to today.
    brand: "spill_crew" | "sampson" | "smart_clean" | "auto" (default auto-detect).
    """
    today = sds_date or date.today().strftime("%d/%m/%Y")
    doc = Document(io.BytesIO(docx_bytes))

    effective_brand = brand if brand and brand != "auto" else _detect_brand(doc)

    if effective_brand == "sampson":
        changes = _rebrand_sampson(doc, today)
    elif effective_brand == "smart_clean":
        changes = _rebrand_smart_clean(doc, today)
    else:
        old_supplier = _detect_supplier_name(doc)
        changes = _apply_supplier_block(doc, today, old_supplier)
        _replace_hyperlink_display_text(doc, old_supplier)
        changes["old_supplier"] = old_supplier

    # Catch any remaining non-CCS emails/URLs in any text run
    sweep = _sweep_email_url(doc)
    changes.setdefault("changes", []).extend(sweep)

    changes["brand"] = effective_brand
    changes["sds_date"] = today

    out_bytes = _save_to_bytes(doc)
    logo_path = _BRAND_LOGOS.get(effective_brand)
    out_bytes = _patch_zip(out_bytes, logo_path)
    return out_bytes, changes


# ---------------------------------------------------------------------------
# Detection
# ---------------------------------------------------------------------------

def _detect_supplier_name(doc: Document) -> str:
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Supplier Name"):
            after = text[len("Supplier Name"):].strip()
            return after
    return ""


# ---------------------------------------------------------------------------
# Supplier block — paragraph-level replacements
# ---------------------------------------------------------------------------

def _apply_supplier_block(doc: Document, sds_date: str, old_supplier: str) -> dict:
    replaced: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()

        # Named supplier fields
        for label, field_key, mode in _SUPPLIER_FIELDS:
            if text.startswith(label):
                new_val = CCS[field_key]
                if mode == "full":
                    _set_full_run(para, f"Supplier Name    {new_val}")
                else:
                    _replace_after_tab(para, new_val)
                replaced.append(f"{label} → {new_val}")
                break

        # SDS Date (may be split across runs: 'SDS', ' ', 'Date', '\t', ...)
        if re.match(r"SDS\s*Date", text, re.IGNORECASE):
            version = _extract_version(text)
            new_date_text = f"{sds_date}, {version}" if version else sds_date
            _replace_after_tab(para, new_date_text)
            replaced.append(f"SDS Date → {new_date_text}")

        # Any body paragraph still containing old supplier name (full or partial)
        if old_supplier and not any(text.startswith(lbl) for lbl, _, _ in _SUPPLIER_FIELDS):
            for term in _supplier_search_terms(old_supplier):
                if term.upper() in text.upper():
                    _replace_text_in_runs(para, term, CCS["supplier_name"])
                    replaced.append(f"Body text: replaced '{term}'")

    return {"changes": replaced}


def _supplier_search_terms(full_name: str) -> list[str]:
    """Return the full name plus a version with legal suffixes stripped."""
    terms = [full_name]
    stripped = re.sub(
        r'\s+(PTY\.?\s*LTD\.?|LIMITED|LTD\.?|INC\.?|LLC|CO\.?|CORP\.?)\s*$',
        '', full_name, flags=re.IGNORECASE,
    ).strip()
    if stripped and stripped.upper() != full_name.upper():
        terms.append(stripped)
    return terms


def _extract_version(text: str) -> str:
    match = re.search(r"Version\s*[\d.]+", text, re.IGNORECASE)
    return match.group(0).strip() if match else ""


# ---------------------------------------------------------------------------
# Brand detection
# ---------------------------------------------------------------------------

def _detect_brand(doc: Document) -> str:
    """Auto-detect supplier brand from document content."""
    texts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    all_text = "\n".join(texts)
    if re.search(r"sampson", all_text, re.IGNORECASE):
        return "sampson"
    if "Chemical Product and Company Identification" in all_text or re.search(r"\bMail Address\b", all_text):
        return "smart_clean"
    return "spill_crew"


# ---------------------------------------------------------------------------
# Post-processing: sweep all text runs for any remaining non-CCS contact info
# ---------------------------------------------------------------------------

def _sweep_email_url(doc: Document) -> list[str]:
    """Replace non-CCS email addresses and website URLs found anywhere in the doc."""
    changes: list[str] = []
    all_paras: list = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    for para in all_paras:
        for run in para.runs:
            text = run.text
            if not text:
                continue
            new_text = re.sub(
                r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}',
                lambda m: CCS["email"] if "compliantcs.com.au" not in m.group(0).lower() else m.group(0),
                text,
            )
            new_text = re.sub(
                r'(?:https?://|www\.)[a-zA-Z0-9.\-/]+',
                lambda m: CCS["website"] if "compliantcs.com.au" not in m.group(0).lower() else m.group(0),
                new_text,
                flags=re.IGNORECASE,
            )
            if new_text != text:
                run.text = new_text
                changes.append(f"Sweep: '{text[:40].strip()}'")
    return changes


# ---------------------------------------------------------------------------
# Sampson Chemical Products handler
# ---------------------------------------------------------------------------

def _replace_cell_lines(cell, lines: list[str]) -> None:
    """Replace paragraph content in a table cell line by line, blanking extras."""
    paras = cell.paragraphs
    for i, text in enumerate(lines):
        if i < len(paras):
            _set_full_run(paras[i], text)
    for i in range(len(lines), len(paras)):
        _set_full_run(paras[i], "")


def _rebrand_sampson(doc: Document, today: str) -> dict:
    """Rebrand a Sampson Chemical Products GHS SDS DOCX."""
    changes: list[str] = []
    ccs_block = [
        CCS["supplier_name"],
        CCS["address"],
        CCS["telephone"],
        CCS["email"],
    ]
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if not cells:
                continue
            label = cells[0].text.strip()
            if label in ("Supplier", "Manufacturer") and len(cells) > 1:
                if re.search(r"sampson|eco\s*pro|bigpond", cells[1].text, re.IGNORECASE):
                    _replace_cell_lines(cells[1], ccs_block)
                    changes.append(f"{label} block → CCS")
            if re.match(r"Revision\s+date", label, re.IGNORECASE) and len(cells) > 1:
                _replace_cell_lines(cells[1], [today])
                changes.append(f"Revision date → {today}")
    for para in doc.paragraphs:
        text = para.text
        if re.search(r"sampson|sampson_office", text, re.IGNORECASE):
            _replace_text_in_runs(para, "Sampson Chemical Products", CCS["supplier_name"])
            _replace_text_in_runs(para, "sampson_office@bigpond.com", CCS["email"])
            changes.append("Body text: replaced Sampson reference")
    return {"changes": changes, "old_supplier": "Sampson Chemical Products"}


# ---------------------------------------------------------------------------
# Smart Clean / Solopak handler
# ---------------------------------------------------------------------------

def _rebrand_smart_clean(doc: Document, today: str) -> dict:
    """Rebrand a Smart Clean / Solopak Australian MSDS DOCX."""
    changes: list[str] = []
    old_supplier = ""
    replacements: dict[str, str] = {
        "Supplier": CCS["supplier_name"],
        "ABN": _CCS_ABN,
        "Mail Address": CCS["address"],
        "Email": CCS["email"],
        "Telephone": CCS["telephone"],
        "Emergency Telephone": f"Poisons Information Centre (National) {CCS['emergency']}",
        "Date of Issue": today,
        "Issue Date": today,
        "Revision Date": today,
        "Prepared By": CCS["supplier_name"],
    }
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            raw_label = row.cells[0].text.strip()
            # Match with or without trailing colon
            label = raw_label.rstrip(":")
            value_cell = row.cells[1]
            if label == "Supplier":
                old_supplier = value_cell.text.strip()
            new_val = replacements.get(label) or replacements.get(raw_label)
            if new_val:
                _replace_cell_lines(value_cell, [new_val])
                changes.append(f"{raw_label} → {new_val}")
    return {"changes": changes, "old_supplier": old_supplier or "Smart Clean / Solopak"}


# ---------------------------------------------------------------------------
# Run manipulation helpers
# ---------------------------------------------------------------------------

def _set_full_run(para, new_text: str) -> None:
    """Replace entire paragraph text via runs (single-run paragraphs)."""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def _replace_after_tab(para, new_value: str) -> None:
    """
    Keep all runs up to and including the first tab character;
    replace everything after with new_value.
    """
    value_started = False
    first_value_run = None

    for run in para.runs:
        if not value_started:
            if "\t" in run.text:
                tab_idx = run.text.index("\t")
                run.text = run.text[: tab_idx + 1]  # keep label + tab
                value_started = True
        else:
            if first_value_run is None:
                first_value_run = run
                run.text = new_value
            else:
                run.text = ""

    if value_started and first_value_run is None:
        para.add_run(new_value)
    elif not value_started:
        # No tab — try to preserve label before colon separator
        full_text = para.text
        colon_m = re.search(r'^(.*?:\s*)', full_text)
        if colon_m:
            _set_full_run(para, colon_m.group(1) + new_value)
        else:
            _set_full_run(para, new_value)


def _replace_text_in_runs(para, old: str, new: str) -> None:
    for run in para.runs:
        if old.upper() in run.text.upper():
            run.text = re.sub(re.escape(old), new, run.text, flags=re.IGNORECASE)


# ---------------------------------------------------------------------------
# Hyperlink display-text update (email / website paragraphs)
# ---------------------------------------------------------------------------

def _replace_hyperlink_display_text(doc: Document, old_supplier: str) -> None:
    """Replace display text of any non-CCS hyperlink with CCS contact details."""
    def _fix(text: str) -> str | None:
        if _CCS_DOMAIN in text.lower():
            return None
        if "@" in text:
            return CCS["email"]
        if re.search(r"https?://|www\.", text, re.IGNORECASE):
            return CCS["website"]
        return None

    for para in doc.paragraphs:
        for hyperlink in para._p.findall(".//" + qn("w:hyperlink")):
            for t_el in hyperlink.findall(".//" + qn("w:t")):
                replacement = _fix(t_el.text or "")
                if replacement:
                    t_el.text = replacement

    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            for hyperlink in para._p.findall(".//" + qn("w:hyperlink")):
                for t_el in hyperlink.findall(".//" + qn("w:t")):
                    replacement = _fix(t_el.text or "")
                    if replacement:
                        t_el.text = replacement


# ---------------------------------------------------------------------------
# ZIP-level patching: logo image + relationship URLs
# ---------------------------------------------------------------------------

def _patch_zip(docx_bytes: bytes, logo_path: Path | None = None) -> bytes:
    """
    Update hyperlink URL targets in .rels files.
    If logo_path is given, also replace the header logo image.
    """
    logo_bytes = logo_path.read_bytes() if logo_path and logo_path.exists() else None
    buf = io.BytesIO()

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin, \
         zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:

        header_image_paths: set[str] = set()
        if logo_bytes:
            for item in zin.infolist():
                if re.search(r"word/_rels/header\d+\.xml\.rels", item.filename, re.IGNORECASE):
                    rels_xml = zin.read(item.filename).decode("utf-8")
                    for m in re.finditer(r'Target="([^"]+)"', rels_xml):
                        target = m.group(1)
                        if re.search(r"\.(png|jpg|jpeg|gif|bmp|tiff|emf|wmf)$", target, re.IGNORECASE):
                            resolved = "word/" + target.lstrip("../")
                            header_image_paths.add(resolved.lower())

        for item in zin.infolist():
            data = zin.read(item.filename)
            fname_lower = item.filename.lower()

            if logo_bytes and fname_lower in header_image_paths:
                data = logo_bytes

            if fname_lower.endswith(".rels"):
                text = data.decode("utf-8")
                text = re.sub(
                    r'(Target=")(mailto:(?![^"]*compliantcs\.com\.au)[^"]+)',
                    lambda m: m.group(1) + _CCS_EMAIL_URL,
                    text, flags=re.IGNORECASE,
                )
                text = re.sub(
                    r'(Target=")(https?://(?![^"]*compliantcs\.com\.au)[^"]+)',
                    lambda m: m.group(1) + _CCS_WEB_URL,
                    text, flags=re.IGNORECASE,
                )
                data = text.encode("utf-8")

            zout.writestr(item, data)

    return buf.getvalue()


# ---------------------------------------------------------------------------
# IO helpers
# ---------------------------------------------------------------------------

def _save_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


import os
import subprocess
import tempfile


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF via LibreOffice headless."""
    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = os.path.join(tmpdir, "input.docx")
        with open(in_path, "wb") as f:
            f.write(docx_bytes)
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, in_path],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(f"PDF conversion failed: {result.stderr.decode()[:300]}")
        out_path = os.path.join(tmpdir, "input.pdf")
        if not os.path.exists(out_path):
            raise RuntimeError("LibreOffice produced no output — is libreoffice installed?")
        with open(out_path, "rb") as f:
            return f.read()
