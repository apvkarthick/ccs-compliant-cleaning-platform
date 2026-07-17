import hashlib
import io
import html
import zipfile
from pathlib import Path

import fitz
from docx import Document

from api.rebrand import CCS, rebrand_sds


def _solid_png(color: int) -> bytes:
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 40, 20), 0)
    pix.clear_with(color)
    return pix.tobytes("png")


def _make_docx_with_header_logo() -> bytes:
    doc = Document()
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.add_run().add_picture(io.BytesIO(_solid_png(0x3366CC)))

    doc.add_paragraph("Supplier Name CLEAN PLUS CHEMICALS PTY LTD")
    doc.add_paragraph("Address 12 OLD STREET")
    doc.add_paragraph("Telephone 1800 000 000")
    doc.add_paragraph("Revision Date: 01/01/2026")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _make_docx_with_body_logo() -> bytes:
    doc = Document()
    doc.add_picture(io.BytesIO(_solid_png(0x3366CC)))
    doc.add_paragraph("Supplier Name CLEAN PLUS CHEMICALS PTY LTD")
    doc.add_paragraph("Address 12 OLD STREET")
    doc.add_paragraph("Telephone 1800 000 000")
    doc.add_paragraph("Revision Date: 01/01/2026")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _make_docx_without_logo() -> bytes:
    doc = Document()
    doc.add_paragraph("Supplier Name SOLOPAK CHEMICALS PTY LTD")
    doc.add_paragraph("Address 12 OLD STREET")
    doc.add_paragraph("Telephone 1800 000 000")
    doc.add_paragraph("Revision Date: 01/01/2026")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _make_docx_with_first_page_header_table() -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    first_header = section.first_page_header
    first_header.paragraphs[0].text = "Solo Pak Pty Ltd"
    table = first_header.add_table(rows=3, cols=2, width=doc.sections[0].page_width)
    table.cell(0, 0).text = "Safety Data Sheet"
    table.cell(0, 1).text = "Page 1"
    table.cell(1, 0).text = "Product"
    table.cell(1, 1).text = "Solopak Test Product"
    table.cell(2, 0).text = ""
    table.cell(2, 1).text = "Issue Date: 1st of July 2026"

    body = doc.add_table(rows=3, cols=2)
    body.cell(0, 0).text = "Supplier"
    body.cell(0, 1).text = "Solopak Test Product"
    body.cell(1, 0).text = "Emergency Telephone:"
    body.cell(1, 1).text = "Poisons Information Centre (National) 000000"
    body.cell(2, 0).text = "Date of Issue"
    body.cell(2, 1).text = "1st of July 2026"

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _asset_digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _embedded_media_digest(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        media_names = sorted(name for name in zf.namelist() if name.startswith("word/media/"))
        assert media_names, "expected at least one embedded media file"
        return hashlib.sha256(zf.read(media_names[0])).hexdigest()


def _embedded_media_digests(docx_bytes: bytes) -> set[str]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        media_names = sorted(name for name in zf.namelist() if name.startswith("word/media/"))
        assert media_names, "expected at least one embedded media file"
        return {hashlib.sha256(zf.read(name)).hexdigest() for name in media_names}


def _header_xml_text(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        parts = [
            zf.read(name).decode("utf-8", errors="ignore")
            for name in zf.namelist()
            if name.startswith("word/header") and name.endswith(".xml")
        ]
    return "\n".join(parts)


def _make_docx_with_body_image_and_messy_address() -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = "Solo Pak Pty Ltd"

    doc.add_picture(io.BytesIO(_solid_png(0x3366CC)))

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Mail Address"
    table.cell(0, 1).text = "PO Box 67, Brisbane86 Crockford Street, NORTHGATE QLD 4013"
    table.cell(1, 0).text = "Supplier"
    table.cell(1, 1).text = "Solo Pak Pty Ltd"

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def test_rebrand_docx_uses_shared_logo_for_cleanplus() -> None:
    src = _make_docx_with_header_logo()
    out_bytes, _summary = rebrand_sds(src, "08/07/2026", brand="cleanplus")

    assert _embedded_media_digest(out_bytes) == _asset_digest(
        Path(r"E:\claude\ccs-compliant-cleaning-platform\api\assets\other-replacement.jpg")
    )


def test_rebrand_docx_uses_solopak_logo_for_solopak() -> None:
    src = _make_docx_with_header_logo()
    out_bytes, _summary = rebrand_sds(src, "08/07/2026", brand="solopak")

    assert _embedded_media_digest(out_bytes) == _asset_digest(
        Path(r"E:\claude\ccs-compliant-cleaning-platform\api\assets\solopak-replacement.jpg")
    )


def test_rebrand_docx_replaces_body_logo_when_no_header_logo_exists() -> None:
    src = _make_docx_with_body_logo()
    out_bytes, _summary = rebrand_sds(src, "08/07/2026", brand="sampson")

    assert _embedded_media_digest(out_bytes) == _asset_digest(
        Path(r"E:\claude\ccs-compliant-cleaning-platform\api\assets\other-replacement.jpg")
    )


def test_rebrand_docx_inserts_solopak_logo_when_source_has_no_logo() -> None:
    src = _make_docx_without_logo()
    out_bytes, summary = rebrand_sds(src, "08/07/2026", brand="solopak")

    assert "Inserted Solopak replacement logo" in summary["changes"]
    assert _embedded_media_digest(out_bytes) == _asset_digest(
        Path(r"E:\claude\ccs-compliant-cleaning-platform\api\assets\solopak-replacement.jpg")
    )


def test_rebrand_docx_inserts_solopak_logo_when_body_image_exists_but_header_is_blank() -> None:
    src = _make_docx_with_body_image_and_messy_address()
    out_bytes, _summary = rebrand_sds(src, "08/07/2026", brand="solopak")

    digests = _embedded_media_digests(out_bytes)
    assert _asset_digest(Path(r"E:\claude\ccs-compliant-cleaning-platform\api\assets\solopak-replacement.jpg")) in digests
    assert hashlib.sha256(_solid_png(0x3366CC)).hexdigest() in digests


def test_rebrand_docx_normalizes_messy_mail_address_cell() -> None:
    src = _make_docx_with_body_image_and_messy_address()
    out_bytes, _summary = rebrand_sds(src, "08/07/2026", brand="solopak")
    out = Document(io.BytesIO(out_bytes))

    assert out.tables[0].cell(0, 1).text == CCS["address"]


def test_rebrand_docx_auto_detects_solopak_from_header_text() -> None:
    src = _make_docx_with_body_image_and_messy_address()
    out_bytes, summary = rebrand_sds(src, "08/07/2026", brand="")
    out = Document(io.BytesIO(out_bytes))

    assert summary["brand"] == "solopak"
    assert "Inserted Solopak replacement logo" in summary["changes"]
    assert _asset_digest(Path(r"E:\claude\ccs-compliant-cleaning-platform\api\assets\solopak-replacement.jpg")) in _embedded_media_digests(out_bytes)
    assert all("Solo Pak Pty Ltd" not in p.text for p in out.sections[0].first_page_header.paragraphs)
    assert any(CCS["supplier_name"] == p.text for p in out.sections[0].first_page_header.paragraphs)


def test_rebrand_docx_updates_first_page_header_issue_date_and_emergency_phone() -> None:
    src = _make_docx_with_first_page_header_table()
    out_bytes, summary = rebrand_sds(src, "08/07/2026", brand="solopak")
    out = Document(io.BytesIO(out_bytes))
    header_xml = _header_xml_text(out_bytes)

    assert any("First page header table updated" in change for change in summary["changes"])
    assert out.sections[0].first_page_header.tables[0].cell(2, 1).text == "Issue Date: 08/07/2026"
    assert out.tables[0].cell(0, 1).text == CCS["supplier_name"]
    assert out.tables[0].cell(1, 1).text == "Poisons Information Centre (National) 131126"
    assert out.tables[0].cell(2, 1).text == "08/07/2026"
    assert "1st of July 2026" not in header_xml
    assert "Issue Date: 08/07/2026" in header_xml
